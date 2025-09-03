import streamlit as st
import re


# ✅ 무조건 첫 Streamlit 명령어
st.set_page_config(
    page_title="PwC 뉴스 분석기",
    page_icon="📊",
    layout="wide",
)



from datetime import datetime, timedelta, timezone
import os
from PIL import Image
import docx
from docx.shared import Pt, RGBColor, Inches
import io
from urllib.parse import urlparse
from googlenews import GoogleNews
import pandas as pd
try:
    import openpyxl
except ImportError:
    st.error("openpyxl 라이브러리가 설치되지 않았습니다. 'pip install openpyxl' 명령어로 설치해주세요.")
    st.stop()
from news_ai import (
    collect_news,
    filter_valid_press,
    filter_excluded_news,
    group_and_select_news,
    evaluate_importance,
    summarize_selected_articles,
    _generate_article_summary,
)

# Import centralized configuration
from config import (
    COMPANY_CATEGORIES,
    COMPANY_KEYWORD_MAP,
    COMPANY_STRUCTURE_NEW,  # 새로운 구조 추가
    COMPANY_STRUCTURE_ENGLISH,  # 영어 키워드 구조 추가
    ANALYSIS_SCOPE_CRITERIA,  # 분석 범위별 특화 기준 추가
    ANALYSIS_SCOPE_SYSTEM_PROMPTS,  # 분석 범위별 시스템 프롬프트 추가
    TRUSTED_PRESS_ALIASES,
    ADDITIONAL_PRESS_ALIASES,
    SYSTEM_PROMPT_1,
    SYSTEM_PROMPT_2,
    SYSTEM_PROMPT_3,
    EXCLUSION_CRITERIA,
    DUPLICATE_HANDLING,
    SELECTION_CRITERIA, 
    GPT_MODELS,
    DEFAULT_GPT_MODEL,
    # 새로 추가되는 회사별 기준들
    COMPANY_ADDITIONAL_EXCLUSION_CRITERIA,
    COMPANY_ADDITIONAL_DUPLICATE_HANDLING,
    COMPANY_ADDITIONAL_SELECTION_CRITERIA
)

# 한국 시간대(KST) 정의
KST = timezone(timedelta(hours=9))


def _clean_html_for_display(text: str) -> str:
    """HTML 태그를 완전히 제거하고 Streamlit 표시용으로 정리"""
    import re
    
    if not text:
        return ""
    
    # 모든 HTML 태그 완전 제거 (더 강력한 정규식)
    text = re.sub(r'<[^<>]*>', '', text)
    text = re.sub(r'<[^<>]*', '', text)  # 닫히지 않은 태그도 제거
    text = re.sub(r'[^<>]*>', '', text)  # 여는 태그가 없는 닫는 태그도 제거
    
    # HTML 엔티티 변환
    html_entities = {
        '&amp;': '&',
        '&lt;': '<',
        '&gt;': '>',
        '&quot;': '"',
        '&#39;': "'",
        '&nbsp;': ' ',
        '&apos;': "'"
    }
    
    for entity, char in html_entities.items():
        text = text.replace(entity, char)
    
    # 특정 HTML 관련 문자열 제거
    text = text.replace('</div>', '')
    text = text.replace('<div>', '')
    text = text.replace('<br>', '\n')
    text = text.replace('<br/>', '\n')
    text = text.replace('<br />', '\n')
    
    # 연속된 공백과 줄바꿈 정리
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()

def _format_ai_summary_for_box(clean_summary, extraction_success):
    """AI 요약을 파란색 박스 안에 포함되도록 HTML 포맷팅 (JSON 형식 지원)"""
    if not clean_summary:
        return ""
    
    if not extraction_success:
        return f"""
            <div class="selection-reason">
                • ⚠️ 원문 추출 실패: {clean_summary}
            </div>
        """
    
    # JSON 형식의 요약이면 그대로 사용 (이미 _format_json_summary에서 HTML로 변환됨)
    # 기존 텍스트 형식이면 기존 방식 사용
    if clean_summary.strip().startswith('<div'):
        # 이미 HTML로 포맷된 JSON 요약
        return f"""
            {clean_summary}
        """
    
    return f"""
            {clean_summary}
    """

def format_date(date_str):
    """Format date to MM/DD format with proper timezone handling"""
    try:
        # Try YYYY-MM-DD format
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%m/%d')
    except Exception:
        try:
            # Try GMT format and convert to KST
            date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
            # Convert UTC to KST (add 9 hours)
            date_obj_kst = date_obj + timedelta(hours=9)
            return date_obj_kst.strftime('%m/%d')
        except Exception:
            try:
                # Try GMT format without timezone indicator
                date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S GMT')
                # Convert UTC to KST (add 9 hours)
                date_obj_kst = date_obj + timedelta(hours=9)
                return date_obj_kst.strftime('%m/%d')
            except Exception:
                # Return original if parsing fails
                return date_str if date_str else '날짜 정보 없음'

# 회사별 추가 기준을 적용하는 함수들
def get_enhanced_exclusion_criteria(companies):
    """회사별 제외 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = EXCLUSION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_duplicate_handling(companies):
    """회사별 중복 처리 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = DUPLICATE_HANDLING
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_selection_criteria(companies):
    """회사별 선택 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = SELECTION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_scope_based_criteria(analysis_scope, criteria_type="selection_criteria", keywords=None):
    """분석 범위에 따른 특화 기준을 반환 (키워드 포함)"""
    if not analysis_scope:
        return ""
    
    # 첫 번째 범위를 주요 기준으로 사용 (본인회사 > 경쟁사 > 산업분야 순서로 우선순위)
    priority_order = ["본인회사", "경쟁사", "산업분야"]
    selected_scope = None
    
    for scope in priority_order:
        if scope in analysis_scope:
            selected_scope = scope
            break
    
    if selected_scope and selected_scope in ANALYSIS_SCOPE_CRITERIA:
        criteria_template = ANALYSIS_SCOPE_CRITERIA[selected_scope].get(criteria_type, "")
        
        # 키워드가 제공된 경우 템플릿에 삽입
        if keywords and "{keywords}" in criteria_template:
            keywords_str = ", ".join(keywords) if isinstance(keywords, list) else str(keywords)
            return criteria_template.format(keywords=keywords_str)
        
        return criteria_template
    
    return ""

def get_scope_based_system_prompts(analysis_scope):
    """분석 범위에 따른 시스템 프롬프트를 반환"""
    if not analysis_scope:
        return SYSTEM_PROMPT_1, SYSTEM_PROMPT_2, SYSTEM_PROMPT_3
    
    # 첫 번째 범위를 주요 기준으로 사용 (본인회사 > 경쟁사 > 산업분야 순서로 우선순위)
    priority_order = ["본인회사", "경쟁사", "산업분야"]
    selected_scope = None
    
    for scope in priority_order:
        if scope in analysis_scope:
            selected_scope = scope
            break
    
    if selected_scope and selected_scope in ANALYSIS_SCOPE_SYSTEM_PROMPTS:
        scope_prompts = ANALYSIS_SCOPE_SYSTEM_PROMPTS[selected_scope]
        return (
            scope_prompts.get("system_prompt_1", SYSTEM_PROMPT_1),
            scope_prompts.get("system_prompt_2", SYSTEM_PROMPT_2),
            scope_prompts.get("system_prompt_3", SYSTEM_PROMPT_3)
        )
    
    return SYSTEM_PROMPT_1, SYSTEM_PROMPT_2, SYSTEM_PROMPT_3
            
# 워드 파일 생성 함수
def create_word_document(keyword, final_selection, analysis=""):
    # 새 워드 문서 생성
    doc = docx.Document()
    
    # 제목 스타일 설정
    title = doc.add_heading(f'PwC 뉴스 분석 보고서: {keyword}', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(208, 74, 2)  # PwC 오렌지 색상
    
    # 분석 요약 추가
    if analysis:
        doc.add_heading('회계법인 관점의 분석 결과', level=1)
        doc.add_paragraph(analysis)
    
    # 선별된 주요 뉴스 추가
    doc.add_heading('선별된 주요 뉴스', level=1)
    
    for i, news in enumerate(final_selection):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {news['title']}").bold = True
        
        # 날짜 정보 추가
        date_str = news.get('date', '날짜 정보 없음')
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"날짜: {date_str}").italic = True
        
        # 선정 사유 추가
        reason = news.get('reason', '')
        if reason:
            doc.add_paragraph(f"선정 사유: {reason}")
        
        # 키워드 추가
        keywords = news.get('keywords', [])
        if keywords:
            doc.add_paragraph(f"키워드: {', '.join(keywords)}")
        
        # 관련 계열사 추가
        affiliates = news.get('affiliates', [])
        if affiliates:
            doc.add_paragraph(f"관련 계열사: {', '.join(affiliates)}")
        
        # 언론사 추가
        press = news.get('press', '알 수 없음')
        doc.add_paragraph(f"언론사: {press}")
        
        # URL 추가
        url = news.get('url', '')
        if url:
            doc.add_paragraph(f"출처: {url}")
        
        # 구분선 추가
        if i < len(final_selection) - 1:
            doc.add_paragraph("").add_run().add_break()
    
    # 날짜 및 푸터 추가
    current_date = datetime.now().strftime("%Y년 %m월 %d일")
    doc.add_paragraph(f"\n보고서 생성일: {current_date}")
    doc.add_paragraph("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")
    
    return doc

# BytesIO 객체로 워드 문서 저장
def get_binary_file_downloader_html(doc, file_name):
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Excel 파일 생성 함수
def create_excel_analysis_report(keyword, final_state, start_date, end_date):
    """전체 뉴스 분석 과정을 Excel로 정리"""
    
    # 데이터 수집
    all_data = []
    
    # 1. 전체 뉴스 데이터 (날짜 필터링 후)
    news_data = final_state.get("news_data", [])
    excluded_news = final_state.get("excluded_news", [])
    borderline_news = final_state.get("borderline_news", [])
    retained_news = final_state.get("retained_news", [])
    grouped_news = final_state.get("grouped_news", [])
    final_selection = final_state.get("final_selection", [])
    not_selected_news = final_state.get("not_selected_news", [])
    
    # 각 뉴스의 상태를 추적하기 위한 딕셔너리
    news_status = {}
    
    # 제외된 뉴스 처리
    for news in excluded_news:
        news_status[news.get('index', -1)] = {
            'status': '제외',
            'reason': news.get('reason', ''),
            'group': '',
            'final_reason': ''
        }
    
    # 보류 뉴스 처리
    for news in borderline_news:
        news_status[news.get('index', -1)] = {
            'status': '보류',
            'reason': news.get('reason', ''),
            'group': '',
            'final_reason': ''
        }
    
    # 유지 뉴스 처리
    for news in retained_news:
        news_status[news.get('index', -1)] = {
            'status': '유지',
            'reason': news.get('reason', ''),
            'group': '',
            'final_reason': ''
        }
    
    # 그룹핑 정보 처리
    for group in grouped_news:
        group_indices = group.get('indices', [])
        selected_index = group.get('selected_index', -1)
        group_reason = group.get('reason', '')
        
        # 그룹 정보를 문자열로 변환
        group_info = f"그룹 {group_indices} (선택: {selected_index})"
        
        for idx in group_indices:
            if idx in news_status:
                news_status[idx]['group'] = group_info
                if idx == selected_index:
                    news_status[idx]['status'] = '그룹 대표 선택'
                else:
                    news_status[idx]['status'] = '그룹 내 미선택'
    
    # 최종 선택된 뉴스 처리
    for news in final_selection:
        # 원본 뉴스에서 인덱스 찾기
        original_index = -1
        for i, original_news in enumerate(news_data, 1):
            if original_news.get('url') == news.get('url') or original_news.get('content') == news.get('title'):
                original_index = i
                break
        
        if original_index in news_status:
            news_status[original_index]['status'] = '최종 선택'
            news_status[original_index]['final_reason'] = news.get('reason', '')
    
    # 최종 선택되지 않은 뉴스 처리
    for news in not_selected_news:
        news_index = news.get('index', -1)
        if news_index in news_status:
            news_status[news_index]['final_reason'] = f"미선택 사유: {news.get('reason', '')}"
    
    # Excel 데이터 생성
    for i, news in enumerate(news_data, 1):
        status_info = news_status.get(i, {
            'status': '상태 불명',
            'reason': '',
            'group': '',
            'final_reason': ''
        })
        
        # 날짜 형식 변환
        date_str = news.get('date', '')
        try:
            if 'GMT' in date_str:
                date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                formatted_date = date_obj.strftime('%Y-%m-%d %H:%M')
            else:
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                formatted_date = date_str
        except:
            formatted_date = date_str if date_str else '날짜 정보 없음'
        
        all_data.append({
            '순번': i,
            '제목': news.get('content', '제목 없음'),
            '언론사': news.get('press', '알 수 없음'),
            '날짜': formatted_date,
            'URL': news.get('url', ''),
            '분석 상태': status_info['status'],
            '1차 분류 사유': status_info['reason'],
            '그룹핑 정보': status_info['group'],
            '최종 선택 사유': status_info['final_reason']
        })
    
    # DataFrame 생성
    df = pd.DataFrame(all_data)
    
    # Excel 파일 생성
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        # 메인 시트
        df.to_excel(writer, sheet_name='전체 뉴스 분석', index=False)
        
        # 워크북과 워크시트 가져오기
        workbook = writer.book
        worksheet = writer.sheets['전체 뉴스 분석']
        
        # 헤더 스타일 설정
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='D04A02', end_color='D04A02', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        # 헤더에 스타일 적용
        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # 열 너비 자동 조정
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            # 최대/최소 너비 설정
            adjusted_width = min(max(max_length + 2, 10), 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # 상태별 색상 구분
        status_colors = {
            '최종 선택': 'C6EFCE',      # 연한 초록
            '그룹 대표 선택': 'C6EFCE',   # 연한 초록
            '제외': 'FFC7CE',           # 연한 빨강
            '보류': 'FFEB9C',           # 연한 노랑
            '유지': 'BDD7EE',           # 연한 파랑
            '그룹 내 미선택': 'F2F2F2'    # 연한 회색
        }
        
        # 데이터 행에 색상 적용
        for row in range(2, len(df) + 2):
            status = worksheet[f'F{row}'].value  # 분석 상태 열
            if status in status_colors:
                fill = PatternFill(start_color=status_colors[status], 
                                 end_color=status_colors[status], 
                                 fill_type='solid')
                for col in range(1, len(df.columns) + 1):
                    worksheet.cell(row=row, column=col).fill = fill
        
        # 요약 시트 추가
        summary_data = [
            ['분석 기간', f"{start_date} ~ {end_date}"],
            ['키워드', keyword],
            ['전체 뉴스 수', len(news_data)],
            ['제외된 뉴스', len(excluded_news)],
            ['보류된 뉴스', len(borderline_news)],
            ['유지된 뉴스', len(retained_news)],
            ['그룹 수', len(grouped_news)],
            ['최종 선택된 뉴스', len(final_selection)]
        ]
        
        summary_df = pd.DataFrame(summary_data, columns=['항목', '값'])
        summary_df.to_excel(writer, sheet_name='분석 요약', index=False)
        
        # 요약 시트 스타일 적용
        summary_ws = writer.sheets['분석 요약']
        for cell in summary_ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # 요약 시트 열 너비 조정
        summary_ws.column_dimensions['A'].width = 20
        summary_ws.column_dimensions['B'].width = 30
    
    bio.seek(0)
    return bio

# 커스텀 CSS
st.markdown("""
<style>
    .title-container {
        display: flex;
        align-items: center;
        gap: 20px;
        margin-bottom: 20px;
    }
    .main-title {
        color: #d04a02;
        font-size: 2.5rem;
        font-weight: 700;
    }
    .news-card {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 15px;
        border-left: 4px solid #d04a02;
    }
    .news-title {
        font-weight: 600;
        font-size: 1.1rem;
    }
    .news-url {
        color: #666;
        font-size: 0.9rem;
    }
    .news-date {
        color: #666;
        font-size: 0.9rem;
        font-style: italic;
        margin-top: 5px;
    }
    .analysis-box {
        background-color: #f5f5ff;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #d04a02;
    }
    .subtitle {
        color: #dc582a;
        font-size: 1.3rem;
        font-weight: 600;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .download-box {
        background-color: #eaf7f0;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #00a36e;
        text-align: center;
    }
    .analysis-section {
        background-color: #f8f9fa;
        border-left: 4px solid #d04a02;
        padding: 20px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .selected-news {
        border-left: 4px solid #0077b6;
        padding: 15px;
        margin: 10px 0;
        background-color: #f0f8ff;
        border-radius: 5px;
    }
    .excluded-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .news-meta {
        color: #666;
        font-size: 0.9em;
        margin: 3px 0;
    }
    .selection-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .keywords {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .affiliates {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0;
        word-break: break-all;
    }
    .news-title-large {
        font-size: 1.2em;
        font-weight: 600;
        color: #000;
        margin-bottom: 8px;
        line-height: 1.4;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0 10px 0;
        word-break: break-all;
    }
    .news-summary {
        color: #444;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .selection-reason {
        color: #666;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .importance-high {
        color: #d04a02;
        font-weight: 700;
        margin: 5px 0;
    }
    .importance-medium {
        color: #0077b6;
        font-weight: 700;
        margin: 5px 0;
    }
    .group-indices {
        color: #666;
        font-size: 0.9em;
    }
    .group-selected {
        color: #00a36e;
        font-weight: 600;
    }
    .group-reason {
        color: #666;
        font-size: 0.9em;
        margin-top: 5px;
    }
    .not-selected-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .importance-low {
        color: #666;
        font-weight: 700;
        margin: 5px 0;
    }
    .not-selected-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .email-preview {
        background-color: white;
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 20px;
        margin: 20px 0;
        overflow-y: auto;
        max-height: 500px;
    }
    .copy-button {
        background-color: #d04a02;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        margin: 10px 0;
    }
    .copy-button:hover {
        background-color: #b33d00;
    }
</style>
""", unsafe_allow_html=True)

# 로고와 제목
col1, col2 = st.columns([1, 5])
with col1:
    # 로고 표시
    logo_path = "pwc_logo.png"
    if os.path.exists(logo_path):
        st.image(logo_path, width=100)
    else:
        st.error("로고 파일을 찾을 수 없습니다. 프로젝트 루트에 'pwc_logo.png' 파일을 추가해주세요.")

with col2:
    st.markdown("<h1 class='main-title'>PwC 뉴스 분석기</h1>", unsafe_allow_html=True)
    st.markdown("회계법인 관점에서 중요한 뉴스를 자동으로 분석하는 AI 도구")

# 기본 선택 카테고리를 Anchor로 설정
COMPANIES = COMPANY_CATEGORIES["Anchor"]

# 세션 상태 초기화 (가장 먼저 실행)
if 'company_keyword_map' not in st.session_state:
    st.session_state.company_keyword_map = COMPANY_KEYWORD_MAP.copy()

# 사이드바 설정
st.sidebar.title("🔍 분석 설정")

# 0단계: 기본 설정
st.sidebar.markdown("### 📋 0단계: 기본 설정")

# 기사 원문 요약 옵션
enable_article_summary = st.sidebar.checkbox(
    "📄 선정된 기사 원문 요약",
    value=False,
    help="선정된 뉴스 기사의 원문을 스크래핑하여 AI 요약을 생성합니다. (시간이 추가로 소요됩니다)"
)

# 유효 언론사 설정 (숨김 처리)
with st.sidebar.expander("📰 유효 언론사 설정 (고급)", expanded=False):
    valid_press_dict = st.text_area(
        "유효 언론사 설정",
        value="""조선일보: ["조선일보", "chosun", "chosun.com"]
    중앙일보: ["중앙일보", "joongang", "joongang.co.kr", "joins.com"]
    동아일보: ["동아일보", "donga", "donga.com"]
    조선비즈: ["조선비즈", "chosunbiz", "biz.chosun.com"]
    매거진한경: ["매거진한경", "magazine.hankyung", "magazine.hankyung.com"]
    한국경제: ["한국경제", "한경", "hankyung", "hankyung.com", "한경닷컴"]
    매일경제: ["매일경제", "매경", "mk", "mk.co.kr"]
    연합뉴스: ["연합뉴스", "yna", "yna.co.kr"]
    파이낸셜뉴스: ["파이낸셜뉴스", "fnnews", "fnnews.com"]
    데일리팜: ["데일리팜", "dailypharm", "dailypharm.com"]
    IT조선: ["it조선", "it.chosun.com", "itchosun"]
    머니투데이: ["머니투데이", "mt", "mt.co.kr"]
    비즈니스포스트: ["비즈니스포스트", "businesspost", "businesspost.co.kr"]
    이데일리: ["이데일리", "edaily", "edaily.co.kr"]
    아시아경제: ["아시아경제", "asiae", "asiae.co.kr"]
    뉴스핌: ["뉴스핌", "newspim", "newspim.com"]
    뉴시스: ["뉴시스", "newsis", "newsis.com"]
    헤럴드경제: ["헤럴드경제", "herald", "heraldcorp", "heraldcorp.com"]""",
        help="분석에 포함할 신뢰할 수 있는 언론사와 그 별칭을 설정하세요. 형식: '언론사: [별칭1, 별칭2, ...]'",
        key="valid_press_dict"
    )



# 구분선 추가
st.sidebar.markdown("---")

# 날짜 필터 설정
st.sidebar.markdown("### 📅 날짜 필터")

# 현재 시간 가져오기
now = datetime.now()

# 기본 시작 날짜/시간 계산
default_start_date = now - timedelta(days=1)

# Set time to 8:00 AM for both start and end - 한국 시간 기준
start_datetime = datetime.combine(default_start_date.date(), 
                                    datetime.strptime("08:00", "%H:%M").time(), KST)
end_datetime = datetime.combine(now.date(), 
                                datetime.strptime("08:00", "%H:%M").time(), KST)

col1, col2 = st.sidebar.columns(2)
with col1:
    start_date = st.date_input(
        "시작 날짜",
        value=default_start_date.date(),
        help="이 날짜부터 뉴스를 검색합니다. 월요일인 경우 지난 금요일, 그 외에는 전일로 자동 설정됩니다."
    )
    start_time = st.time_input(
        "시작 시간",
        value=start_datetime.time(),
        help="시작 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )
with col2:
    end_date = st.date_input(
        "종료 날짜",
        value=now.date(),
        help="이 날짜까지의 뉴스를 검색합니다."
    )
    end_time = st.time_input(
        "종료 시간",
        value=end_datetime.time(),
        help="종료 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )

# 구분선 추가
st.sidebar.markdown("---")

# 1단계: 제외 판단 기준

# 기업 선택 섹션 제목
st.sidebar.markdown("### 🏢 분석할 기업 선택")

# 현대차그룹 분석 구조 사용
st.sidebar.markdown("#### 🔥 현대차그룹 분석 구조")

# 현대차그룹 선택 (현재는 하나만 있음)
selected_group = st.sidebar.selectbox(
    "그룹 선택",
    options=list(COMPANY_STRUCTURE_NEW.keys()),
    index=0,
    help="분석할 기업 그룹을 선택하세요"
)

# 분석 범위 선택
analysis_scope = st.sidebar.multiselect(
    "📊 분석 범위 선택",
    options=["본인회사", "경쟁사", "산업분야"],
    default=["산업분야"],
    help="어떤 범위의 뉴스를 분석할지 선택하세요"
)

# 선택된 키워드들을 수집
selected_keywords = []

if "본인회사" in analysis_scope:
    st.sidebar.markdown("##### 🏠 본인회사")
    own_company_types = st.sidebar.multiselect(
        "계열사 유형 선택",
        options=list(COMPANY_STRUCTURE_NEW[selected_group]["본인회사"].keys()),
        default=["핵심계열사"],
        help="포함할 계열사 유형을 선택하세요"
    )
    
    for comp_type in own_company_types:
        companies = COMPANY_STRUCTURE_NEW[selected_group]["본인회사"][comp_type]
        selected_companies_in_type = st.sidebar.multiselect(
            f"{comp_type}",
            options=companies,
            default=companies[:3] if len(companies) > 3 else companies,
            key=f"own_{comp_type}"
        )
        selected_keywords.extend(selected_companies_in_type)

if "경쟁사" in analysis_scope:
    st.sidebar.markdown("##### ⚔️ 경쟁사")
    competitor_types = st.sidebar.multiselect(
        "경쟁사 유형 선택",
        options=list(COMPANY_STRUCTURE_NEW[selected_group]["경쟁사"].keys()),
        default=["국내경쟁사"],
        help="분석할 경쟁사 유형을 선택하세요"
    )
    
    for comp_type in competitor_types:
        companies = COMPANY_STRUCTURE_NEW[selected_group]["경쟁사"][comp_type]
        selected_companies_in_type = st.sidebar.multiselect(
            f"{comp_type}",
            options=companies,
            default=companies[:2] if len(companies) > 2 else companies,
            key=f"comp_{comp_type}"
        )
        selected_keywords.extend(selected_companies_in_type)

if "산업분야" in analysis_scope:
    st.sidebar.markdown("##### 🏭 산업분야")
    industry_fields = st.sidebar.multiselect(
        "산업 분야 선택",
        options=list(COMPANY_STRUCTURE_NEW[selected_group]["산업분야"].keys()),
        default=["배터리_Cell_Module_System","배터리_Charging_Simulation","전동화_Motor_Drive","전동화_Control_Electronics","내연기관_연소","연료_대체연료","배출가스_후처리","구동계_변속기","하이브리드_열관리","시뮬레이션_제어","기타"],
        help="분석할 산업 분야를 선택하세요. 각 분야별로 세부 키워드를 선택할 수 있습니다."
    )

    # 각 산업분야별로 세부 키워드 선택 가능하게 함
    for field in industry_fields:
        with st.sidebar.expander(f"🔧 {field} 세부 키워드 설정"):
            # 새로운 영어 키워드 구조에서 키워드 가져오기
            field_keywords = COMPANY_STRUCTURE_NEW[selected_group]["산업분야"][field]
            selected_field_keywords = st.multiselect(
                f"🇺🇸 {field} 키워드",
                options=field_keywords,
                default=field_keywords,  # 기본적으로 모든 키워드 선택
                key=f"field_{field}"
            )
            
            # 선택된 키워드들을 세션 상태에 저장
            if selected_field_keywords:
                st.session_state.company_keyword_map[field] = selected_field_keywords
                selected_keywords.append(field)  # 분야 이름을 키워드로 추가
            else:
                # 아무것도 선택하지 않은 경우 기본값으로 분야 이름만 사용
                st.session_state.company_keyword_map[field] = [field]
                selected_keywords.append(field)
            
            # 선택된 키워드 요약 표시
            if selected_field_keywords:
                st.info(f"선택됨: {len(selected_field_keywords)}개 키워드")

# 최종 선택된 키워드들을 companies로 설정
selected_companies = selected_keywords[:10]  # 최대 10개 제한

# 키워드 맵 업데이트 - 이미 위에서 처리됨 (산업분야는 사용자 선택으로, 일반 키워드는 자기 자신)
for keyword in selected_companies:
    if 'company_keyword_map' not in st.session_state:
            st.session_state.company_keyword_map = {}
        
    # 산업분야가 아닌 일반 키워드는 자기 자신만 포함
    if keyword not in COMPANY_STRUCTURE_NEW[selected_group]["산업분야"]:
        st.session_state.company_keyword_map[keyword] = [keyword]





# 선택된 키워드 미리보기
st.sidebar.markdown("### 🔍 선택된 키워드 미리보기")
st.sidebar.markdown(f"총 **{len(selected_companies)}개** 키워드가 선택되었습니다.")

# 세션 상태는 이미 위에서 초기화됨

# 추가 키워드 설정
st.sidebar.markdown("#### ➕ 추가 키워드 설정")
additional_keywords_text = st.sidebar.text_area(
    "직접 키워드 추가",
    value="",
    placeholder="키워드1, 키워드2, 키워드3\n(쉼표로 구분하여 입력)",
    help="위에서 선택한 키워드 외에 추가로 분석하고 싶은 키워드를 입력하세요.",
    key="additional_keywords"
)

# 추가 키워드 처리
additional_keywords = []
if additional_keywords_text.strip():
    # 쉼표로 구분된 키워드들을 파싱
    additional_keywords = [kw.strip() for kw in additional_keywords_text.split(',') if kw.strip()]
    
    # 추가 키워드들을 세션 상태에 추가
    for keyword in additional_keywords:
        if keyword not in st.session_state.company_keyword_map:
            st.session_state.company_keyword_map[keyword] = [keyword]

# 최종 키워드 리스트 업데이트 (선택된 키워드 + 추가 키워드)
final_selected_companies = selected_companies + additional_keywords
final_selected_companies = list(dict.fromkeys(final_selected_companies))  # 중복 제거하면서 순서 유지
final_selected_companies = final_selected_companies[:15]  # 최대 15개로 확장

# 미리보기 버튼 - 모든 검색어 확인
with st.sidebar.expander("🔍 전체 검색 키워드 미리보기"):
    # 선택된 키워드들과 실제 검색어들을 표시
    st.markdown("**📊 선택된 분석 대상**")
    
    if len(selected_companies) > 0:
        st.markdown("**🏢 구조 기반 선택:**")
        for i, keyword in enumerate(selected_companies, 1):
            st.markdown(f"**{i}. {keyword}**")
            # 실제 검색에 사용될 키워드들 표시
            if 'company_keyword_map' in st.session_state:
                search_terms = st.session_state.company_keyword_map.get(keyword, [keyword])
                if len(search_terms) > 1:
                    st.write(f"   🇺🇸 영어 키워드: {', '.join(search_terms)}")
                    st.write(f"   📊 총 {len(search_terms)}개 키워드")
                else:
                    st.write(f"   🔍 검색어: {search_terms[0]}")
    
    if len(additional_keywords) > 0:
        st.markdown("**➕ 직접 추가:**")
        for i, keyword in enumerate(additional_keywords, len(selected_companies) + 1):
            st.markdown(f"**{i}. {keyword}**")
            st.write(f"   🔍 검색어: {keyword}")
    
    if len(final_selected_companies) == 0:
        st.info("분석 대상을 선택해주세요.")
    else:
        st.success(f"총 **{len(final_selected_companies)}개** 키워드로 분석합니다.")

# 선택된 키워드들을 통합 (검색용) - 최종 리스트 사용
keywords = final_selected_companies

# 중복 제거는 이미 위에서 처리됨
# keywords = list(set(keywords))

# 구분선 추가
st.sidebar.markdown("---")



# 구분선 추가
st.sidebar.markdown("---")

# GPT 모델 선택 섹션
st.sidebar.markdown("### 🤖 GPT 모델 선택")

selected_model = st.sidebar.selectbox(
    "분석에 사용할 GPT 모델을 선택하세요",
    options=list(GPT_MODELS.keys()),
    index=list(GPT_MODELS.keys()).index(DEFAULT_GPT_MODEL) if DEFAULT_GPT_MODEL in GPT_MODELS else 0,
    format_func=lambda x: f"{x} - {GPT_MODELS[x]}",
    help="각 모델의 특성:\n" + "\n".join([f"• {k}: {v}" for k, v in GPT_MODELS.items()])
)

# 모델 설명 표시
st.sidebar.markdown(f"""
<div style='background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin-bottom: 20px;'>
    <strong>선택된 모델:</strong> {selected_model}<br>
    <strong>특징:</strong> {GPT_MODELS[selected_model]}
</div>
""", unsafe_allow_html=True)

# 구분선 추가
st.sidebar.markdown("---")

# 검색 결과 수 - 고정 값으로 설정
max_results = 100

# AI 프롬프트 및 분석 기준 설정
st.sidebar.markdown("### 🤖 AI 분석 단계별 설정")

# 프롬프트 미리보기용 분석 범위 선택
prompt_preview_scope = st.sidebar.selectbox(
    "분석 범위 선택",
    options=["기본", "본인회사", "경쟁사", "산업분야"],
    index=0,
    help="어떤 분석 범위의 AI 프롬프트와 기준을 미리보기할지 선택하세요"
)

# 선택된 범위에 따른 프롬프트와 기준 가져오기
if prompt_preview_scope == "기본":
    preview_system_prompt_1 = SYSTEM_PROMPT_1
    preview_system_prompt_2 = SYSTEM_PROMPT_2
    preview_system_prompt_3 = SYSTEM_PROMPT_3
    preview_exclusion_criteria = EXCLUSION_CRITERIA
    preview_selection_criteria = SELECTION_CRITERIA
else:
    # 시스템 프롬프트
    preview_prompts = get_scope_based_system_prompts([prompt_preview_scope])
    preview_system_prompt_1 = preview_prompts[0]
    preview_system_prompt_2 = preview_prompts[1] 
    preview_system_prompt_3 = preview_prompts[2]
    
    # 분석 기준
    dummy_keywords = ["키워드1", "키워드2"]
    preview_exclusion_criteria = EXCLUSION_CRITERIA + "\n\n" + get_scope_based_criteria([prompt_preview_scope], "exclusion_criteria", dummy_keywords)
    preview_selection_criteria = get_scope_based_criteria([prompt_preview_scope], "selection_criteria", dummy_keywords)
    if not preview_selection_criteria:
        preview_selection_criteria = SELECTION_CRITERIA

# 단계별 설정
st.sidebar.markdown(f"#### 📋 1단계: 제외 판단 ({prompt_preview_scope})")

# 1단계 시스템 프롬프트
system_prompt_1 = st.sidebar.text_area(
    "🤖 시스템 프롬프트",
    value=preview_system_prompt_1,
    help=f"{prompt_preview_scope} 범위의 1단계 제외 판단에 사용되는 AI 시스템 프롬프트입니다.",
    key=f"system_prompt_1_{prompt_preview_scope}",
    height=200
)

# 1단계 분석 기준 (유저 프롬프트)
exclusion_criteria = st.sidebar.text_area(
    "👤 분석 기준 (유저 프롬프트)",
    value=preview_exclusion_criteria,
    help=f"{prompt_preview_scope} 범위의 제외 기준입니다. 실제 분석 시 선택된 키워드가 적용됩니다.",
    key=f"exclusion_criteria_{prompt_preview_scope}",
    height=200
)

st.sidebar.markdown("---")
st.sidebar.markdown(f"#### 📋 2단계: 그룹핑 ({prompt_preview_scope})")

# 2단계 시스템 프롬프트
system_prompt_2 = st.sidebar.text_area(
    "🤖 시스템 프롬프트",
    value=preview_system_prompt_2,
    help=f"{prompt_preview_scope} 범위의 2단계 그룹핑에 사용되는 AI 시스템 프롬프트입니다.",
    key=f"system_prompt_2_{prompt_preview_scope}",
    height=200
)

# 2단계 분석 기준 (유저 프롬프트)
duplicate_handling = st.sidebar.text_area(
    "👤 분석 기준 (유저 프롬프트)",
    value=DUPLICATE_HANDLING,
    help="중복된 뉴스를 처리하는 기준을 설정하세요.",
    key=f"duplicate_handling_{prompt_preview_scope}",
    height=200
)

st.sidebar.markdown("---")
st.sidebar.markdown(f"#### 📋 3단계: 중요도 평가 ({prompt_preview_scope})")

# 3단계 시스템 프롬프트
system_prompt_3 = st.sidebar.text_area(
    "🤖 시스템 프롬프트",
    value=preview_system_prompt_3,
    help=f"{prompt_preview_scope} 범위의 3단계 중요도 평가에 사용되는 AI 시스템 프롬프트입니다.",
    key=f"system_prompt_3_{prompt_preview_scope}",
    height=200
)

# 3단계 분석 기준 (유저 프롬프트)
selection_criteria = st.sidebar.text_area(
    "👤 분석 기준 (유저 프롬프트)",
    value=preview_selection_criteria,
    help=f"{prompt_preview_scope} 범위의 선택 기준입니다. 실제 분석 시 선택된 키워드가 적용됩니다.",
    key=f"selection_criteria_{prompt_preview_scope}",
    height=200
)

# 응답 형식 설정
response_format = st.sidebar.text_area(
    "📝 응답 형식",
    value="""선택된 뉴스 인덱스: [1, 3, 5]와 같은 형식으로 알려주세요.

각 선택된 뉴스에 대해:
제목: (뉴스 제목)
언론사: (언론사명)
발행일: (발행일자)
선정 사유: (구체적인 선정 이유)
분석 키워드: (해당 기업 그룹의 주요 계열사들)

[제외된 주요 뉴스]
제외된 중요 뉴스들에 대해:
인덱스: (뉴스 인덱스)
제목: (뉴스 제목)
제외 사유: (구체적인 제외 이유)""",
    help="분석 결과의 출력 형식을 설정하세요.",
    key="response_format",
    height=200
)

# 최종 프롬프트 생성
analysis_prompt = f"""
당신은 회계법인의 전문 애널리스트입니다. 아래 뉴스 목록을 분석하여 회계법인 관점에서 가장 중요한 뉴스를 선별하세요. 

[선택 기준]
{selection_criteria}

[제외 대상]
{exclusion_criteria}

[응답 요구사항]
1. 선택 기준에 부합하는 뉴스가 많다면 최대 3개까지 선택 가능합니다.
2. 선택 기준에 부합하는 뉴스가 없다면, 그 이유를 명확히 설명해주세요.

[응답 형식]
다음과 같은 JSON 형식으로 응답해주세요:

{{
    "selected_news": [
        {{
            "index": 1,
            "title": "뉴스 제목",
            "press": "언론사명",
            "date": "발행일자",
            "reason": "선정 사유",
            "keywords": ["키워드1", "키워드2"]
        }},
        ...
    ],
    "excluded_news": [
        {{
            "index": 2,
            "title": "뉴스 제목",
            "reason": "제외 사유"
        }},
        ...
    ]
}}

[유효 언론사]
{valid_press_dict}

[중복 처리 기준]
{duplicate_handling}
"""

# 메인 컨텐츠
if st.button("뉴스 분석 시작", type="primary"):
    # 이메일 미리보기를 위한 전체 내용 저장
    email_content = "[Client Intelligence]\n\n"
    
    # 모든 키워드 분석 결과를 저장할 딕셔너리
    all_results = {}
    
    # 모든 키워드의 전체 분석 상태를 저장할 딕셔너리 (Excel 통합용)
    all_analysis_states = {}
    
    for i, company in enumerate(final_selected_companies, 1):
        with st.spinner(f"'{company}' 관련 뉴스를 수집하고 분석 중입니다..."):
            # 해당 회사의 연관 키워드 확장 (세션 상태에서 가져옴)
            company_keywords = st.session_state.company_keyword_map.get(company, [company])
            
            # 연관 키워드 표시
            st.write(f"'{company}' 연관 키워드로 검색 중: {', '.join(company_keywords)}")
            
            # 범위별 특화 기준 적용
            if 'analysis_scope' in locals():
                # 범위별 특화 기준 적용 (실제 검색 키워드 포함)
                scope_selection_criteria = get_scope_based_criteria(analysis_scope, "selection_criteria", company_keywords)
                scope_exclusion_criteria = get_scope_based_criteria(analysis_scope, "exclusion_criteria", company_keywords)
                
                # 범위별 시스템 프롬프트 적용
                scope_system_prompt_1, scope_system_prompt_2, scope_system_prompt_3 = get_scope_based_system_prompts(analysis_scope)
                
                # 범위별 기준이 있으면 사용, 없으면 기본 기준 사용
                if scope_selection_criteria:
                    enhanced_selection_criteria = scope_selection_criteria
                else:
                    enhanced_selection_criteria = selection_criteria
                    
                if scope_exclusion_criteria:
                    enhanced_exclusion_criteria = exclusion_criteria + "\n\n" + scope_exclusion_criteria
                else:
                    enhanced_exclusion_criteria = exclusion_criteria
                    
                # 중복 처리는 기본 기준 사용
                enhanced_duplicate_handling = duplicate_handling
                
                # 현재 선택된 분석 범위 확인
                priority_order = ["본인회사", "경쟁사", "산업분야"]
                selected_scope = None
                for scope in priority_order:
                    if scope in analysis_scope:
                        selected_scope = scope
                        break
                
                st.info(f"🎯 분석 범위별 특화 기준 적용: {', '.join(analysis_scope)}")
                st.info(f"🔍 실제 검색 키워드: {', '.join(company_keywords)}")
                if selected_scope:
                    st.info(f"🤖 시스템 프롬프트 변경: {selected_scope} 분석 모드")
                
            else:
                # analysis_scope가 없는 경우
                scope_system_prompt_1, scope_system_prompt_2, scope_system_prompt_3 = SYSTEM_PROMPT_1, SYSTEM_PROMPT_2, SYSTEM_PROMPT_3
                
                base_exclusion = exclusion_criteria
                base_duplicate = duplicate_handling
                base_selection = selection_criteria
                
                # 해당 회사의 추가 특화 기준만 가져오기 (세션 상태에서)
                # 세션 상태가 초기화되지 않은 경우를 위한 안전장치
                if 'company_additional_exclusion_criteria' not in st.session_state:
                    st.session_state.company_additional_exclusion_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.copy()
                if 'company_additional_duplicate_handling' not in st.session_state:
                    st.session_state.company_additional_duplicate_handling = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.copy()
                if 'company_additional_selection_criteria' not in st.session_state:
                    st.session_state.company_additional_selection_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.copy()
                    
                company_additional_exclusion = st.session_state.company_additional_exclusion_criteria.get(company, "")
                company_additional_duplicate = st.session_state.company_additional_duplicate_handling.get(company, "")
                company_additional_selection = st.session_state.company_additional_selection_criteria.get(company, "")
                
                # 사용자 수정 기준 + 해당 회사 특화 기준 결합
                enhanced_exclusion_criteria = base_exclusion + company_additional_exclusion
                enhanced_duplicate_handling = base_duplicate + company_additional_duplicate  
                enhanced_selection_criteria = base_selection + company_additional_selection
            
            # initial_state 설정 부분 직전에 valid_press_dict를 딕셔너리로 변환하는 코드 추가
            # 텍스트 에어리어의 내용을 딕셔너리로 변환
            valid_press_config = {}
            try:
                # 문자열에서 딕셔너리 파싱
                lines = valid_press_dict.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if line and ': ' in line:
                        press_name, aliases_str = line.split(':', 1)
                        try:
                            # 문자열 형태의 리스트를 실제 리스트로 변환
                            aliases = eval(aliases_str.strip())
                            valid_press_config[press_name.strip()] = aliases
                            print(f"[DEBUG] Valid press 파싱 성공: {press_name.strip()} -> {aliases}")
                        except Exception as e:
                            print(f"[DEBUG] Valid press 파싱 실패: {line}, 오류: {str(e)}")
            except Exception as e:
                print(f"[DEBUG] Valid press 전체 파싱 실패: {str(e)}")
                # 오류 발생 시 빈 딕셔너리 사용
                valid_press_config = {}
            
            print(f"[DEBUG] 파싱된 valid_press_dict: {valid_press_config}")
            

            
            # 각 키워드별 상태 초기화
            initial_state = {
                "news_data": [], 
                "filtered_news": [], 
                "analysis": "", 
                "keyword": company_keywords,  # 회사별 확장 키워드 리스트 전달
                "model": selected_model,
                "excluded_news": [],
                "borderline_news": [],
                "retained_news": [],
                "grouped_news": [],
                "final_selection": [],
                # 회사별 enhanced 기준들 적용
                "exclusion_criteria": enhanced_exclusion_criteria,
                "duplicate_handling": enhanced_duplicate_handling,
                "selection_criteria": enhanced_selection_criteria,
                "system_prompt_1": scope_system_prompt_1,
                "user_prompt_1": "",
                "llm_response_1": "",
                "system_prompt_2": scope_system_prompt_2,
                "user_prompt_2": "",
                "llm_response_2": "",
                "system_prompt_3": scope_system_prompt_3,
                "user_prompt_3": "",
                "llm_response_3": "",
                "not_selected_news": [],
                "original_news_data": [],
                # 언론사 설정 추가 (파싱된 딕셔너리 사용)
                "valid_press_dict": valid_press_config,
                # 날짜 필터 정보 추가
                "start_datetime": datetime.combine(start_date, start_time, KST),
                "end_datetime": datetime.combine(end_date, end_time, KST)
                #"start_datetime": start_datetime,
                #"end_datetime": end_datetime
            }
            
            
            print(f"[DEBUG] start_datetime: {datetime.combine(start_date, start_time)}")
            print(f"[DEBUG] end_datetime: {datetime.combine(end_date, end_time)}")
            
            # 1단계: 뉴스 수집
            st.write("1단계: 뉴스 수집 중...")
            state_after_collection = collect_news(initial_state)
            
            # 2단계: 유효 언론사 필터링 (글로벌 뉴스를 위해 건너뜀)
            st.write("2단계: 날짜 필터링 중...")
            state_after_press_filter = state_after_collection  # 필터링 없이 그대로 전달
            
            # 3단계: 제외 판단
            st.write("3단계: 제외 판단 중...")
            state_after_exclusion = filter_excluded_news(state_after_press_filter)
            
            # 4단계: 그룹핑
            st.write("4단계: 그룹핑 중...")
            state_after_grouping = group_and_select_news(state_after_exclusion)
            
            # 5단계: 중요도 평가
            st.write("5단계: 중요도 평가 중...")
            final_state = evaluate_importance(state_after_grouping)
            
            # 6단계: 기사 원문 요약 (옵션)
            if enable_article_summary and final_state.get("final_selection"):
                st.write("6단계: 선정된 기사 원문 요약 중...")
                final_state = summarize_selected_articles(final_state)

            # 선정된 뉴스가 없는 경우 메시지 표시
            if len(final_state["final_selection"]) == 0:
                st.info("선정 기준에 부합하는 뉴스가 없습니다.")

            # 키워드별 분석 결과 저장
            all_results[company] = final_state["final_selection"]
            
            # 전체 분석 상태 저장 (Excel 통합용)
            all_analysis_states[company] = final_state
            
            # 키워드 구분선 추가
            st.markdown("---")
            
            # 키워드별 섹션 구분
            st.markdown(f"## 📊 {company} 분석 결과")
            

            
            # 전체 뉴스 표시 (필터링 전)
            with st.expander(f"📰 '{company}' 관련 전체 뉴스 (필터링 전)"):
                for i, news in enumerate(final_state.get("original_news_data", []), 1):
                    date_str = news.get('date', '날짜 정보 없음')
                    url = news.get('url', 'URL 정보 없음')
                    press = news.get('press', '알 수 없음')
                    st.markdown(f"""
                    <div class="news-card">
                        <div class="news-title">{i}. {news['content']}</div>
                        <div class="news-meta">📰 {press}</div>
                        <div class="news-date">📅 {date_str}</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 유효 언론사 필터링된 뉴스 표시
            with st.expander(f"📰 '{company}' 관련 유효 언론사 뉴스"):
                for i, news in enumerate(final_state["news_data"]):
                    date_str = news.get('date', '날짜 정보 없음')
                    url = news.get('url', 'URL 정보 없음')
                    press = news.get('press', '알 수 없음')
                    st.markdown(f"""
                    <div class="news-card">
                        <div class="news-title">{i+1}. {news['content']}</div>
                        <div class="news-meta">📰 {press}</div>
                        <div class="news-date">📅 {date_str}</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 2단계: 유효 언론사 필터링 결과 표시
            st.markdown("<div class='subtitle'>🔍 2단계: 유효 언론사 필터링 결과</div>", unsafe_allow_html=True)
            st.markdown(f"유효 언론사 뉴스: {len(final_state['news_data'])}개")
            
            # 3단계: 제외/보류/유지 뉴스 표시
            st.markdown("<div class='subtitle'>🔍 3단계: 뉴스 분류 결과</div>", unsafe_allow_html=True)
            
            # 제외된 뉴스
            with st.expander("❌ 제외된 뉴스"):
                for news in final_state["excluded_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 보류 뉴스
            with st.expander("⚠️ 보류 뉴스"):
                for news in final_state["borderline_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 유지 뉴스
            with st.expander("✅ 유지 뉴스"):
                for news in final_state["retained_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 4단계: 그룹핑 결과 표시
            st.markdown("<div class='subtitle'>🔍 4단계: 뉴스 그룹핑 결과</div>", unsafe_allow_html=True)
            
            with st.expander("📋 그룹핑 결과 보기"):
                for group in final_state["grouped_news"]:
                    st.markdown(f"""
                    <div class="analysis-section">
                        <h4>그룹 {group['indices']}</h4>
                        <p>선택된 기사: {group['selected_index']}</p>
                        <p>선정 이유: {group['reason']}</p>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 5단계: 최종 선택 결과 표시
            st.markdown("<div class='subtitle'>🔍 5단계: 최종 선택 결과</div>", unsafe_allow_html=True)
            
            # 최종 선정된 뉴스가 있는 경우에만 표시
            if final_state["final_selection"]:
                st.markdown("### 📰 최종 선정된 뉴스")  
                news_style = ""
                reason_prefix = "선별 이유: "
            
            # 최종 선정된 뉴스 표시
            for news in final_state["final_selection"]:
                # 날짜 형식 변환
                
                date_str = format_date(news.get('date', ''))
                
                try:
                    # YYYY-MM-DD 형식으로 가정
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        # GMT 형식 시도
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'

                url = news.get('url', 'URL 정보 없음')
                press = news.get('press', '언론사 정보 없음')
                
                # AI 요약 텍스트 준비 (HTML 태그 완전 제거)
                clean_summary = None
                if news.get('ai_summary'):
                    clean_summary = _clean_html_for_display(news['ai_summary'])
                
                # 뉴스 정보 표시
                st.markdown(f"""
                    <div class="selected-news" style="{news_style}">
                        <div class="news-title-large">{news['title']} ({formatted_date})</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                        <div class="selection-reason">
                            • {reason_prefix}{news['reason']}
                        </div>
                        <div class="news-summary">
                            • 키워드: {', '.join(news['keywords'])} | 관련 계열사: {', '.join(news['affiliates'])} | 언론사: {press}
                        </div>
                        {_format_ai_summary_for_box(clean_summary, news.get('extraction_success', False))}
                """, unsafe_allow_html=True)
                
                # 구분선 추가
                st.markdown("---")
            
            # 선정되지 않은 뉴스 표시
            if final_state.get("not_selected_news"):
                with st.expander("❌ 선정되지 않은 뉴스"):
                    for news in final_state["not_selected_news"]:
                        st.markdown(f"""
                        <div class="not-selected-news">
                            <div class="news-title">{news['index']}. {news['title']}</div>
                            <div class="importance-low">💡 중요도: {news['importance']}</div>
                            <div class="not-selected-reason">❌ 미선정 사유: {news['reason']}</div>
                        </div>
                        """, unsafe_allow_html=True)
            
            # 디버그 정보
            with st.expander("디버그 정보"):
                st.markdown("### 1단계: 제외 판단")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_1", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_1", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_1", "없음"))
                
                st.markdown("### 2단계: 그룹핑")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_2", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_2", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_2", "없음"))
                
                st.markdown("### 3단계: 중요도 평가")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_3", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_3", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_3", "없음"))
                

            
            # 이메일 내용 추가
            email_content += f"{i}. {company}\n"
            for news in final_state["final_selection"]:
                # 날짜 형식 변환
                date_str = news.get('date', '')
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'
                
                url = news.get('url', '')
                email_content += f"  - {news['title']} ({formatted_date}) {url}\n"
            email_content += "\n"
            
            # 키워드 구분선 추가
            st.markdown("---")

    # 모든 키워드 분석이 끝난 후 통합 Excel 다운로드
    st.markdown("---")
    st.markdown("### 📊 전체 분석 결과 Excel 다운로드")
    
    # 통합 Excel 생성 함수
    def create_integrated_excel_report(all_results_data, start_date, end_date):
        """모든 키워드의 분석 결과를 하나의 시트에 통합한 Excel 파일 생성"""
        bio = io.BytesIO()
        
        # 모든 키워드의 데이터를 통합할 리스트
        integrated_data = []
        
        # 각 키워드별 데이터 처리
        for company, final_state in all_results_data.items():
            if not final_state:  # final_state가 비어있으면 건너뛰기
                continue
                
            # 해당 키워드의 데이터 처리
            news_data = final_state.get("news_data", [])
            excluded_news = final_state.get("excluded_news", [])
            borderline_news = final_state.get("borderline_news", [])
            retained_news = final_state.get("retained_news", [])
            grouped_news = final_state.get("grouped_news", [])
            final_selection = final_state.get("final_selection", [])
            not_selected_news = final_state.get("not_selected_news", [])
            
            # 뉴스 상태 추적
            news_status = {}
            
            # 제외된 뉴스 처리
            for news in excluded_news:
                news_status[news.get('index', -1)] = {
                    'status': '제외',
                    'reason': news.get('reason', ''),
                    'group': '',
                    'final_reason': ''
                }
            
            # 보류 뉴스 처리
            for news in borderline_news:
                news_status[news.get('index', -1)] = {
                    'status': '보류',
                    'reason': news.get('reason', ''),
                    'group': '',
                    'final_reason': ''
                }
            
            # 유지 뉴스 처리
            for news in retained_news:
                news_status[news.get('index', -1)] = {
                    'status': '유지',
                    'reason': news.get('reason', ''),
                    'group': '',
                    'final_reason': ''
                }
            
            # 그룹핑 정보 처리
            for group in grouped_news:
                group_indices = group.get('indices', [])
                selected_index = group.get('selected_index', -1)
                group_info = f"그룹 {group_indices} (선택: {selected_index})"
                
                for idx in group_indices:
                    if idx in news_status:
                        news_status[idx]['group'] = group_info
                        if idx == selected_index:
                            news_status[idx]['status'] = '그룹 대표 선택'
                        else:
                            news_status[idx]['status'] = '그룹 내 미선택'
            
            # 최종 선택된 뉴스 처리
            for news in final_selection:
                original_index = -1
                for i, original_news in enumerate(news_data, 1):
                    if original_news.get('url') == news.get('url') or original_news.get('content') == news.get('title'):
                        original_index = i
                        break
                
                if original_index in news_status:
                    news_status[original_index]['status'] = '최종 선택'
                    news_status[original_index]['final_reason'] = news.get('reason', '')
            
            # 최종 선택되지 않은 뉴스 처리
            for news in not_selected_news:
                news_index = news.get('index', -1)
                if news_index in news_status:
                    news_status[news_index]['final_reason'] = f"미선택 사유: {news.get('reason', '')}"
            
            # 해당 키워드의 모든 뉴스 데이터를 통합 리스트에 추가
            for i, news in enumerate(news_data, 1):
                status_info = news_status.get(i, {
                    'status': '상태 불명',
                    'reason': '',
                    'group': '',
                    'final_reason': ''
                })
                
                # 날짜 형식 변환
                date_str = news.get('date', '')
                try:
                    if 'GMT' in date_str:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%Y-%m-%d %H:%M')
                    else:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        formatted_date = date_str
                except:
                    formatted_date = date_str if date_str else '날짜 정보 없음'
                
                # AI 요약 정보 파싱 (최종 선택된 뉴스인 경우에만)
                ai_title_korean = ""
                ai_summary_oneline = ""
                ai_details = ""
                extraction_success = False
                
                if status_info['status'] == '최종 선택':
                    # final_selection에서 해당 뉴스 찾기
                    for selected_news in final_selection:
                        if (selected_news.get('url') == news.get('url') or 
                            selected_news.get('title') == news.get('content')):
                            
                            ai_summary = selected_news.get('ai_summary', '')
                            extraction_success = selected_news.get('extraction_success', False)
                            
                            if ai_summary and extraction_success:
                                try:
                                    import json
                                    import re
                                    
                                    # JSON 응답에서 코드 블록 제거
                                    json_text = ai_summary.strip()
                                    if json_text.startswith("```json"):
                                        json_text = json_text[7:]
                                    if json_text.startswith("```"):
                                        json_text = "\n".join(json_text.split("\n")[1:])
                                    if json_text.endswith("```"):
                                        json_text = "\n".join(json_text.split("\n")[:-1])
                                    
                                    json_text = json_text.strip()
                                    summary_data = json.loads(json_text)
                                    
                                    ai_title_korean = summary_data.get('title_korean', summary_data.get('title', ''))
                                    ai_summary_oneline = summary_data.get('summary_oneline', summary_data.get('summary', ''))
                                    details = summary_data.get('details', [])
                                    
                                    # 세부 내용을 문자열로 결합
                                    if details:
                                        ai_details = " | ".join(details)
                                    
                                except Exception as e:
                                    print(f"AI 요약 파싱 실패: {e}")
                                    # JSON 파싱 실패 시 원본 텍스트 사용
                                    ai_summary_oneline = ai_summary[:200] + "..." if len(ai_summary) > 200 else ai_summary
                            elif ai_summary:
                                # 추출 실패했지만 오류 메시지가 있는 경우
                                ai_summary_oneline = ai_summary
                            break
                
                integrated_data.append({
                    '키워드': company,  # 맨 앞 컬럼에 키워드 추가
                    '순번': i,
                    '제목': news.get('content', '제목 없음'),
                    '언론사': news.get('press', '알 수 없음'),
                    '날짜': formatted_date,
                    'URL': news.get('url', ''),
                    '분석 상태': status_info['status'],
                    '1차 분류 사유': status_info['reason'],
                    '그룹핑 정보': status_info['group'],
                    '최종 선택 사유': status_info['final_reason'],
                    'AI 번역 제목': ai_title_korean,
                    'AI 핵심 요약': ai_summary_oneline,
                    'AI 세부 내용': ai_details,
                    '원문 추출 성공': '성공' if extraction_success else '실패' if status_info['status'] == '최종 선택' else ''
                })
        
        # DataFrame 생성
        df = pd.DataFrame(integrated_data)
        
        with pd.ExcelWriter(bio, engine='openpyxl') as writer:
            # 통합 데이터를 하나의 시트에 저장
            df.to_excel(writer, sheet_name='전체 뉴스 분석', index=False)
            
            # 스타일 적용
            worksheet = writer.sheets['전체 뉴스 분석']
            from openpyxl.styles import Font, PatternFill, Alignment
            
            header_font = Font(bold=True, color='FFFFFF')
            header_fill = PatternFill(start_color='D04A02', end_color='D04A02', fill_type='solid')
            header_alignment = Alignment(horizontal='center', vertical='center')
            
            # 헤더에 스타일 적용
            for cell in worksheet[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            
            # 열 너비 자동 조정
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = min(max(max_length + 2, 10), 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # 상태별 색상 구분
            status_colors = {
                '최종 선택': 'C6EFCE',
                '그룹 대표 선택': 'C6EFCE',
                '제외': 'FFC7CE',
                '보류': 'FFEB9C',
                '유지': 'BDD7EE',
                '그룹 내 미선택': 'F2F2F2'
            }
            
            # 데이터 행에 색상 적용 (분석 상태 컬럼 찾기)
            status_col_index = None
            for col_idx, col_name in enumerate(df.columns, 1):
                if col_name == '분석 상태':
                    status_col_index = col_idx
                    break
            
            if status_col_index:
                status_col_letter = chr(64 + status_col_index)  # A=65, B=66, ...
                for row in range(2, len(df) + 2):
                    status = worksheet[f'{status_col_letter}{row}'].value
                    if status in status_colors:
                        fill = PatternFill(start_color=status_colors[status], 
                                         end_color=status_colors[status], 
                                         fill_type='solid')
                        for col in range(1, len(df.columns) + 1):
                            worksheet.cell(row=row, column=col).fill = fill
            
            # 요약 시트 추가
            summary_data = []
            summary_data.append(['분석 기간', f"{start_date} ~ {end_date}"])
            summary_data.append(['분석 키워드 수', len(all_results_data)])
            summary_data.append(['전체 뉴스 수', len(integrated_data)])
            summary_data.append(['', ''])  # 빈 행
            
            # 키워드별 통계
            for company, final_state in all_results_data.items():
                if final_state:
                    news_count = len(final_state.get("news_data", []))
                    selected_count = len(final_state.get("final_selection", []))
                    summary_data.append([f"{company} - 전체 뉴스", news_count])
                    summary_data.append([f"{company} - 최종 선택", selected_count])
            
            summary_df = pd.DataFrame(summary_data, columns=['항목', '값'])
            summary_df.to_excel(writer, sheet_name='분석 요약', index=False)
            
            # 요약 시트 스타일 적용
            summary_ws = writer.sheets['분석 요약']
            for cell in summary_ws[1]:
                cell.font = Font(bold=True, color='FFFFFF')
                cell.fill = PatternFill(start_color='D04A02', end_color='D04A02', fill_type='solid')
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            summary_ws.column_dimensions['A'].width = 25
            summary_ws.column_dimensions['B'].width = 15
        
        bio.seek(0)
        return bio
    
    # 통합 Excel 다운로드 버튼
    if any(all_results.values()):  # 분석 결과가 있는 경우에만 표시
        try:
            # 통합 Excel 파일 생성
            integrated_excel = create_integrated_excel_report(
                all_results_data=all_analysis_states,
                start_date=start_date.strftime('%Y-%m-%d'),
                end_date=end_date.strftime('%Y-%m-%d')
            )
            
            current_time = datetime.now().strftime("%Y%m%d_%H%M")
            filename = f"뉴스분석_통합결과_{current_time}.xlsx"
            
            st.download_button(
                label="📊 전체 분석 결과 Excel 다운로드",
                data=integrated_excel,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                help="모든 키워드의 전체 뉴스 분석 과정을 통합한 Excel 파일을 다운로드합니다.",
                type="primary"
            )
            
            st.success(f"✅ {len(all_analysis_states)}개 키워드의 통합 분석 결과를 Excel로 다운로드할 수 있습니다!")
            
        except Exception as e:
            st.error(f"통합 Excel 파일 생성 중 오류가 발생했습니다: {str(e)}")
    
    st.markdown("---")
    st.markdown("### 📧 이메일용 HTML 요약")
    st.markdown("아래 HTML을 복사하여 이메일로 전송하실 수 있습니다.")
    
    # 모든 키워드의 최종 선정 기사들을 하나의 리스트로 통합
    all_final_news = []
    for company, results in all_results.items():
        for news in results:
            news['source_keyword'] = company
            all_final_news.append(news)
    
    if all_final_news:
        with st.spinner("선정된 기사들의 상세 요약을 생성하는 중..."):
            # 웹 스크래퍼 초기화
            from web_scraper import HybridNewsWebScraper
            scraper = HybridNewsWebScraper(
                openai_api_key=os.getenv('OPENAI_API_KEY'),
                enable_ai_fallback=True
            )
            
            # HTML 이메일 내용 생성
            html_content = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; color: #333; }
        .header { background-color: #d04a02; color: white; padding: 20px; text-align: center; }
        .article { border: 1px solid #ddd; margin: 20px 0; padding: 20px; border-radius: 8px; }
        .article-title { font-size: 1.3em; font-weight: bold; color: #d04a02; margin-bottom: 10px; }
        .article-meta { color: #666; font-size: 0.9em; margin-bottom: 15px; }
        .korean-title { font-size: 1.2em; font-weight: bold; color: #333; margin: 15px 0 10px 0; }
        .oneline-summary { background-color: #f0f8ff; padding: 12px; border-radius: 6px; margin: 15px 0; border-left: 4px solid #0077b6; }
        .details { margin: 15px 0; }
        .details li { margin-bottom: 8px; line-height: 1.4; }
        .original-url { margin-top: 15px; }
        .original-url a { color: #0077b6; text-decoration: none; }
        .original-url a:hover { text-decoration: underline; }
        .footer { background-color: #f8f9fa; padding: 15px; text-align: center; color: #666; margin-top: 30px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>PwC 뉴스 분석 보고서</h1>
        <p>생성일: """ + datetime.now().strftime("%Y년 %m월 %d일 %H:%M") + """</p>
    </div>
"""
            
            # 각 기사별로 상세 정보 추가
            for i, news in enumerate(all_final_news, 1):
                url = news.get('url', '')
                title = news.get('title', '제목 없음')
                press = news.get('press', '알 수 없음')
                date_str = news.get('date', '')
                source_keyword = news.get('source_keyword', '')
                reason = news.get('reason', '')
                
                # 날짜 형식 변환
                try:
                    if 'GMT' in date_str:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%Y-%m-%d')
                    else:
                        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                        formatted_date = date_str
                except:
                    formatted_date = date_str if date_str else '날짜 정보 없음'
                
                # Google News URL 디코딩
                original_url = url
                if 'news.google.com' in url:
                    decoded_url = scraper._resolve_google_news_url_simple(url, timeout=10)
                    if decoded_url:
                        original_url = decoded_url
                
                # 기사 원문 요약 처리 (옵션이 활성화된 경우에만)
                summary_html = ""
                
                if enable_article_summary and news.get('ai_summary'):
                    # 이미 생성된 AI 요약 사용
                    summary = news.get('ai_summary', '')
                    
                    # JSON 파싱하여 각 요소 추출
                    try:
                        import json
                        import re
                        
                        # JSON 응답에서 코드 블록 제거
                        json_text = summary.strip()
                        if "<div" in json_text:  # 이미 HTML로 포맷된 경우
                            summary_html = json_text
                        else:
                            # JSON 파싱 시도
                            if json_text.startswith("```json"):
                                json_text = json_text[7:]
                            if json_text.startswith("```"):
                                json_text = "\n".join(json_text.split("\n")[1:])
                            if json_text.endswith("```"):
                                json_text = "\n".join(json_text.split("\n")[:-1])
                            
                            json_text = json_text.strip()
                            summary_data = json.loads(json_text)
                            
                            korean_title = summary_data.get('title_korean', '번역 제목 없음')
                            oneline_summary = summary_data.get('summary_oneline', '요약 없음')
                            details = summary_data.get('details', [])
                            
                            # 세부 내용 HTML 생성
                            details_html = ""
                            if details:
                                details_html = "<ul class='details'>"
                                for detail in details:
                                    details_html += f"<li>{detail}</li>"
                                details_html += "</ul>"
                            
                            summary_html = f"""
                            <div class="korean-title">{korean_title}</div>
                            <div class="oneline-summary"><strong>핵심 요약:</strong> {oneline_summary}</div>
                            {details_html}
                            """
                    except:
                        # JSON 파싱 실패 시 원본 HTML 사용
                        if summary:
                            summary_html = summary
                        else:
                            summary_html = "<div>요약 파싱 실패</div>"
                
                elif enable_article_summary:
                    # 원문 요약 옵션이 켜져있지만 요약이 없는 경우
                    summary_html = "<div style='color: #666; font-style: italic;'>원문 요약을 생성하지 못했습니다.</div>"
                else:
                    # 원문 요약 옵션이 꺼져있는 경우
                    summary_html = "<div style='color: #666; font-style: italic;'>원문 요약이 비활성화되었습니다.</div>"
                
                # HTML에 기사 정보 추가
                html_content += f"""
    <div class="article">
        <div class="article-title">{i}. {title}</div>
        <div class="article-meta">
            <strong>날짜:</strong> {formatted_date} | 
            <strong>언론사:</strong> {press} | 
            <strong>키워드:</strong> {source_keyword}
        </div>
        <div class="article-meta">
            <strong>선정 이유:</strong> {reason}
        </div>
        
        {summary_html}
        
        <div class="original-url">
            <strong>원문 링크:</strong> <a href="{original_url}" target="_blank">{original_url}</a>
        </div>
    </div>
"""
                
                # 진행상황 표시
                st.write(f"기사 {i}/{len(all_final_news)} 처리 완료: {title[:50]}...")
            
            # HTML 마무리
            html_content += """
    <div class="footer">
        <p>© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구</p>
    </div>
</body>
</html>
"""
            
            # HTML 내용 표시
            st.markdown("#### 📋 생성된 HTML 이메일")
            st.text_area(
                "HTML 코드 (복사하여 사용하세요)",
                value=html_content,
                height=400,
                help="이 HTML 코드를 복사하여 이메일 본문에 붙여넣으세요."
            )
            
            # HTML 미리보기
            st.markdown("#### 👀 이메일 미리보기")
            st.components.v1.html(html_content, height=600, scrolling=True)
            
            st.success(f"🎉 총 {len(all_final_news)}개 기사의 HTML 이메일이 생성되었습니다!")
    
    else:
        st.info("선정된 기사가 없어 HTML 이메일을 생성할 수 없습니다.")
    




else:
    # 초기 화면 설명 (주석 처리됨)
    """
    ### 👋 PwC 뉴스 분석기에 오신 것을 환영합니다!
    
    이 도구는 입력한 키워드에 대한 최신 뉴스를 자동으로 수집하고, 회계법인 관점에서 중요한 뉴스를 선별하여 분석해드립니다.
    
    #### 주요 기능:
    1. 최신 뉴스 자동 수집 (기본 100개)
    2. 신뢰할 수 있는 언론사 필터링
    3. 6단계 AI 기반 뉴스 분석 프로세스:
       - 1단계: 뉴스 수집 - 키워드 기반으로 최신 뉴스 데이터 수집
       - 2단계: 유효 언론사 필터링 - 신뢰할 수 있는 언론사 선별
       - 3단계: 제외/보류/유지 판단 - 회계법인 관점에서의 중요도 1차 분류
       - 4단계: 유사 뉴스 그룹핑 - 중복 기사 제거 및 대표 기사 선정
       - 5단계: 중요도 평가 및 최종 선정 - 회계법인 관점의 중요도 평가
       - 6단계: 필요시 재평가 - 선정된 뉴스가 없을 경우 AI가 기준을 완화하여 재평가
    4. 선별된 뉴스에 대한 상세 정보 제공
       - 제목 및 날짜
       - 원문 링크
       - 선별 이유
       - 키워드, 관련 계열사, 언론사 정보
    5. 분석 결과 이메일 형식 미리보기
    
    #### 사용 방법:
    1. 사이드바에서 분석할 기업을 선택하세요 (최대 10개)
       - 기본 제공 기업 목록에서 선택
       - 새로운 기업 직접 추가 가능
    2. GPT 모델을 선택하세요
       - gpt-4o: 빠르고 실시간 (기본값)
    3. 날짜 필터를 설정하세요
       - 기본값: 어제 또는 지난 금요일(월요일인 경우)부터 오늘까지
    4. "뉴스 분석 시작" 버튼을 클릭하세요
    
    #### 분석 결과 확인:
    - 각 키워드별 최종 선정된 중요 뉴스
    - 선정 과정의 중간 결과(제외/보류/유지, 그룹핑 등)
    - 선정된 모든 뉴스의 요약 이메일 미리보기
    - 디버그 정보 (시스템 프롬프트, AI 응답 등)
    
    """

# 푸터
st.markdown("---")
st.markdown("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")